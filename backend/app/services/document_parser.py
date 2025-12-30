"""
文档解析服务

功能：
1. 支持 docx, md, pdf 格式文档解析
2. 智能多层次切片策略
3. 图片提取和原位置插入
"""

import hashlib
import os
import re
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Dict, List

# ==================== 数据结构 ====================


@dataclass
class DocumentChunk:
    """文档块数据结构"""

    chunk_id: str  # 块唯一ID
    text: str  # 文本内容
    source_file: str  # 来源文件
    section: str  # 所属章节
    page_number: int  # 页码（如有）
    images: List[str]  # 关联图片路径列表
    metadata: Dict[str, Any]  # 其他元数据

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return asdict(self)


# ==================== 配置类 ====================


@dataclass
class ChunkingConfig:
    """切片配置参数"""

    # 基础切片参数
    max_chunk_size: int = 800  # 最大块大小（字符数）
    min_chunk_size: int = 100  # 最小块大小（避免过小的块）
    chunk_overlap: int = 100  # 块之间的重叠字符数

    # 切片策略
    split_by_title: bool = True  # 是否按标题切分
    split_by_paragraph: bool = True  # 是否按段落切分
    force_max_size: bool = True  # 是否强制限制最大长度

    # 标题识别模式
    title_patterns: List[str] = None  # 额外的标题识别正则模式

    # 图片处理
    distribute_images_evenly: bool = True  # 是否均匀分配图片到各chunk

    def __post_init__(self):
        """初始化默认值"""
        if self.title_patterns is None:
            # 中文标题模式：一、二、1. 2. 第一章 等
            self.title_patterns = [
                r"^[一二三四五六七八九十]+[、\.]\s*.+",  # 一、 或 一.
                r"^\d+[、\.]\s*.+",  # 1、 或 1.
                r"^第[一二三四五六七八九十\d]+[章节部分]\s*.+",  # 第一章
                r"^[（\(]\d+[）\)]\s*.+",  # (1) 或 （1）
            ]


# ==================== 优化的文档解析器 ====================


class OptimizedDocumentParser:
    """
    优化的文档解析器 - 支持智能切片

    主要优化：
    1. 多层次切片：标题 > 段落 > 固定长度
    2. 长度控制：保证 chunk 在合理范围
    3. 重叠机制：提高检索召回率
    4. 增强标题识别：支持中文标题模式
    """

    def __init__(
        self,
        image_output_dir: str = "./extracted_images",
        config: ChunkingConfig = None,
    ):
        """
        初始化解析器

        Args:
            image_output_dir: 图片提取输出目录
            config: 切片配置
        """
        self.image_output_dir = Path(image_output_dir)
        self.image_output_dir.mkdir(parents=True, exist_ok=True)
        self.config = config or ChunkingConfig()

    def parse(self, file_path: str) -> List[DocumentChunk]:
        """
        解析文档，返回带图片关联的文本块

        Args:
            file_path: 文档路径

        Returns:
            DocumentChunk 列表
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return self._parse_docx(file_path)
        elif suffix == ".md":
            return self._parse_markdown(file_path)
        elif suffix == ".pdf":
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    def parse_directory(
        self, dir_path: str, extensions: List[str] = None
    ) -> List[DocumentChunk]:
        """
        解析目录下的所有文档

        Args:
            dir_path: 目录路径
            extensions: 要处理的文件扩展名列表，默认 ['.docx', '.md']

        Returns:
            所有文档的 DocumentChunk 列表
        """
        if extensions is None:
            extensions = [".docx", ".md"]

        dir_path = Path(dir_path)
        all_chunks = []

        for ext in extensions:
            for file_path in dir_path.glob(f"*{ext}"):
                try:
                    chunks = self.parse(str(file_path))
                    all_chunks.extend(chunks)
                    print(f"✅ 解析完成: {file_path.name} ({len(chunks)} 个块)")
                except Exception as e:
                    print(f"❌ 解析失败: {file_path.name} - {e}")

        return all_chunks

    # ==================== Word 文档解析 ====================

    def _parse_docx(self, file_path: Path) -> List[DocumentChunk]:
        """
        解析 Word 文档

        简化逻辑：直接提取所有文本内容，只对图片做定位
        """
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(str(file_path))
        doc_name = file_path.stem

        # 构建元素列表，包含文本和图片的位置信息
        elements = []
        image_index = 0

        # 遍历段落，提取文本和图片
        for para in doc.paragraphs:
            # 检查段落样式，判断是否为标题
            style_name = para.style.name if para.style else ""
            is_heading = style_name.startswith("Heading") or style_name.startswith(
                "标题"
            )

            # 检查段落中的图片
            para_images = []
            for run in para.runs:
                drawing_elements = run._element.findall(
                    ".//a:blip",
                    {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
                )
                for drawing in drawing_elements:
                    embed_id = drawing.get(qn("r:embed"))
                    if embed_id:
                        try:
                            image_part = doc.part.related_parts[embed_id]
                            image_ext = image_part.content_type.split("/")[-1]
                            if image_ext == "jpeg":
                                image_ext = "jpg"
                            image_index += 1
                            image_name = f"{doc_name}_image{image_index}.{image_ext}"
                            output_path = self.image_output_dir / image_name

                            with open(output_path, "wb") as f:
                                f.write(image_part.blob)

                            para_images.append(str(output_path))
                        except Exception as e:
                            print(f"  ⚠️  提取图片失败: {e}")

            # 添加元素（直接提取所有文本）
            text = para.text.strip()
            if text or para_images:
                elements.append(
                    {
                        "type": "heading" if is_heading else "paragraph",
                        "text": text,
                        "images": para_images,
                    }
                )

        # 处理表格中的内容
        table_text_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_text_count += 1
                        elements.append(
                            {
                                "type": "paragraph",
                                "text": cell_text,
                                "images": [],
                            }
                        )

        print(
            f"  📊 python-docx 解析: {len(elements)} 个元素，{image_index} 张图片，{table_text_count} 个表格文本"
        )

        # 使用优化的切片策略
        return self._docx_elements_to_chunks(
            elements=elements,
            source_file=str(file_path),
        )

    def _docx_elements_to_chunks(
        self,
        elements: List[Dict],
        source_file: str,
    ) -> List[DocumentChunk]:
        """
        将 python-docx 解析的元素转换为 chunks

        简化逻辑：直接按标题分组，图片在原位置插入占位符
        """
        chunks = []
        current_section = "文档开始"

        # 按标题分组
        sections = []
        section_data = {
            "section": current_section,
            "texts": [],
            "images": [],
            "page": 0,
        }

        for elem in elements:
            elem_type = elem["type"]
            text = elem["text"]
            images = elem.get("images", [])

            # 检查是否为标题
            is_title = elem_type == "heading"

            # 额外检查：中文标题模式
            if (
                not is_title
                and text
                and self.config.split_by_title
                and self._is_custom_title(text)
            ):
                is_title = True

            if is_title and text:
                # 保存之前的 section
                if section_data["texts"] or section_data["images"]:
                    sections.append(section_data)

                # 新 section
                current_section = text
                section_data = {
                    "section": current_section,
                    "texts": [],
                    "images": [],
                    "page": 0,
                }
            else:
                # 添加文本到当前 section
                if text:
                    section_data["texts"].append(text)
                # 在文本中插入图片占位符
                if images:
                    for img_path in images:
                        placeholder = f"[IMG:{img_path}]"
                        section_data["texts"].append(placeholder)
                        section_data["images"].append(img_path)
                        print(f"  🖼️  图片已定位: {os.path.basename(img_path)}")

        # 保存最后一个 section
        if section_data["texts"] or section_data["images"]:
            sections.append(section_data)

        # 如果没有识别到任何 section，创建默认 section
        if not sections:
            all_texts = [e["text"] for e in elements if e["text"]]
            all_images = []
            for e in elements:
                all_images.extend(e.get("images", []))
            sections = [
                {
                    "section": "文档内容",
                    "texts": all_texts,
                    "images": all_images,
                    "page": 0,
                }
            ]

        total_images = sum(len(s["images"]) for s in sections)
        print(f"  📊 识别到 {len(sections)} 个章节，{total_images} 张图片已定位")

        # 对每个 section 应用长度限制和重叠
        for section_info in sections:
            section_text = "\n".join(section_info["texts"])
            section_images = section_info.get("images", [])

            # 如果 section 过长，进行切分
            if len(section_text) > self.config.max_chunk_size:
                print(
                    f"  ✂️  章节 [{section_info['section']}] 过长 ({len(section_text)} 字符)，进行切分..."
                )
                sub_chunks = self._split_text_with_overlap(
                    text=section_text,
                    source_file=source_file,
                    section=section_info["section"],
                    images=section_images,
                    page_number=section_info["page"],
                )
                chunks.extend(sub_chunks)
            else:
                # 直接创建 chunk
                if section_text.strip():
                    chunk = self._create_chunk(
                        text=section_text,
                        source_file=source_file,
                        section=section_info["section"],
                        images=section_images,
                        page_number=section_info["page"],
                    )
                    chunks.append(chunk)

        print(f"  ✅ 生成 {len(chunks)} 个优化的文档块")
        return chunks

    def _extract_docx_images(self, file_path: Path) -> List[str]:
        """从 Word 文档中提取嵌入图片（备用方法）"""
        extracted = []
        doc_name = file_path.stem

        try:
            with zipfile.ZipFile(file_path, "r") as z:
                for file_info in z.namelist():
                    if file_info.startswith("word/media/"):
                        # 提取图片
                        image_data = z.read(file_info)
                        image_name = f"{doc_name}_{Path(file_info).name}"
                        output_path = self.image_output_dir / image_name

                        with open(output_path, "wb") as f:
                            f.write(image_data)

                        extracted.append(str(output_path))
        except Exception as e:
            print(f"⚠️ 提取图片失败: {e}")

        return extracted

    # ==================== Markdown 解析 ====================

    def _parse_markdown(self, file_path: Path) -> List[DocumentChunk]:
        """解析 Markdown 文档"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # 提取所有图片引用
        image_pattern = r"!\[([^\]]*)\]\(([^)]+)\)"
        images_in_doc = re.findall(image_pattern, content)

        # 构建图片路径映射
        doc_dir = file_path.parent
        image_map = {}

        for alt_text, img_path in images_in_doc:
            full_path = doc_dir / img_path
            if full_path.exists():
                image_map[img_path] = str(full_path)

        # 使用优化的切分策略
        chunks = self._split_markdown_optimized(content, str(file_path), image_map)

        return chunks

    def _split_markdown_optimized(
        self, content: str, source_file: str, image_map: Dict[str, str]
    ) -> List[DocumentChunk]:
        """优化的 Markdown 切分"""
        # 第一步：保护代码块，避免被标题分割打断
        code_block_pattern = r"(```[\s\S]*?```)"
        code_blocks = []

        def protect_code_block(match):
            """将代码块替换为占位符"""
            code_blocks.append(match.group(0))
            return f"__CODE_BLOCK_{len(code_blocks) - 1}__"

        # 保护代码块
        protected_content = re.sub(code_block_pattern, protect_code_block, content)

        # 按标题分割（支持 1-3 级标题）
        section_pattern = r"(^#{1,3}\s+.+$)"
        parts = re.split(section_pattern, protected_content, flags=re.MULTILINE)

        chunks = []
        current_section = "文档开始"
        current_text = []
        current_images = []

        def restore_code_blocks(text: str) -> str:
            """恢复代码块"""
            result = text
            for i, code_block in enumerate(code_blocks):
                result = result.replace(f"__CODE_BLOCK_{i}__", code_block)
            return result

        for part in parts:
            part = part.strip()
            if not part:
                continue

            # 检查是否是标题
            if re.match(r"^#{1,3}\s+", part):
                # 保存之前的块（使用优化切分）
                if current_text:
                    text = restore_code_blocks("\n".join(current_text))
                    sub_chunks = self._split_text_with_overlap(
                        text=text,
                        source_file=source_file,
                        section=current_section,
                        images=current_images.copy(),
                    )
                    chunks.extend(sub_chunks)
                    current_text = []
                    current_images = []

                current_section = re.sub(r"^#+\s*", "", part)
            else:
                # 先恢复代码块再提取图片
                restored_part = restore_code_blocks(part)

                # 提取该部分中的图片
                img_refs = re.findall(r"!\[[^\]]*\]\(([^)]+)\)", restored_part)
                for img_ref in img_refs:
                    if (
                        img_ref in image_map
                        and image_map[img_ref] not in current_images
                    ):
                        current_images.append(image_map[img_ref])

                # 清理图片引用为占位符，但保留代码块
                clean_text = re.sub(
                    r"!\[([^\]]*)\]\([^)]+\)", r"[图片: \1]", restored_part
                )
                if clean_text.strip():
                    current_text.append(clean_text)

        # 保存最后一个块
        if current_text:
            text = restore_code_blocks("\n".join(current_text))
            sub_chunks = self._split_text_with_overlap(
                text=text,
                source_file=source_file,
                section=current_section,
                images=current_images.copy(),
            )
            chunks.extend(sub_chunks)

        return chunks

    # ==================== PDF 解析 ====================

    def _parse_pdf(self, file_path: Path) -> List[DocumentChunk]:
        """解析 PDF 文档"""
        try:
            from unstructured.partition.pdf import partition_pdf
        except ImportError:
            raise ImportError('请安装 PDF 支持: pip install "unstructured[pdf]"')

        elements = partition_pdf(
            filename=str(file_path),
            strategy="hi_res",
            extract_images_in_pdf=True,
            extract_image_block_types=["Image", "Table"],
            extract_image_block_output_dir=str(self.image_output_dir),
        )

        extracted_images = list(
            self.image_output_dir.glob(f"{file_path.stem}*.png")
        ) + list(self.image_output_dir.glob(f"{file_path.stem}*.jpg"))

        return self._elements_to_chunks_optimized(
            elements=elements,
            source_file=str(file_path),
            extracted_images=[str(p) for p in extracted_images],
        )

    # ==================== 优化的切片核心算法 ====================

    def _elements_to_chunks_optimized(
        self,
        elements,
        source_file: str,
        extracted_images: List[str],
    ) -> List[DocumentChunk]:
        """
        优化的元素转换方法 - 支持图片原位置插入

        策略：
        1. 按 Title/Header 初步分组
        2. 识别自定义标题模式（中文标题）
        3. 识别 Image 元素，在原位置插入 [IMG:path] 占位符
        4. 对每组应用长度限制和重叠
        """
        chunks = []
        current_section = "文档开始"
        current_page = 0
        image_index = 0  # 图片索引，用于匹配提取的图片

        # 第一遍：按标题初步分组，同时处理图片占位符
        sections = []
        section_data = {
            "section": current_section,
            "texts": [],
            "page": 0,
            "images": [],  # 该 section 包含的图片路径
        }

        for element in elements:
            element_type = element.category
            text = element.text.strip() if element.text else ""

            # 检查是否是图片元素
            if element_type == "Image":
                # 在文本中插入图片占位符
                if image_index < len(extracted_images):
                    img_path = extracted_images[image_index]
                    placeholder = f"[IMG:{img_path}]"
                    section_data["texts"].append(placeholder)
                    section_data["images"].append(img_path)
                    image_index += 1
                    print(
                        f"  🖼️  发现图片元素，插入占位符: {os.path.basename(img_path)}"
                    )
                continue

            # 标题识别
            is_title = element_type in ["Title", "Header"]

            # 额外检查：中文标题模式
            if (
                not is_title
                and text
                and self.config.split_by_title
                and self._is_custom_title(text)
            ):
                is_title = True

            if is_title and text:
                # 保存之前的section
                if section_data["texts"]:
                    sections.append(section_data)

                # 新section
                current_section = text
                current_page = self._get_page_number(element)
                section_data = {
                    "section": current_section,
                    "texts": [],
                    "page": current_page,
                    "images": [],
                }

            elif text:
                # 添加文本到当前section
                section_data["texts"].append(text)

        # 保存最后一个section
        if section_data["texts"]:
            sections.append(section_data)

        # 如果没有识别到任何section，创建默认section
        if not sections:
            all_texts = [
                e.text.strip() for e in elements if e.text and e.category != "Image"
            ]
            sections = [
                {
                    "section": "文档内容",
                    "texts": all_texts,
                    "page": 0,
                    "images": extracted_images,
                }
            ]

        # 处理未被分配的图片（如果 unstructured 没有识别出 Image 元素）
        remaining_images = (
            extracted_images[image_index:]
            if image_index < len(extracted_images)
            else []
        )
        if remaining_images:
            print(
                f"  ⚠️  有 {len(remaining_images)} 张图片未能定位到原始位置，将均匀分配"
            )

        print(f"  📊 识别到 {len(sections)} 个章节，{image_index} 张图片已定位")

        # 第二遍：对每个section应用长度限制和重叠
        for section_info in sections:
            section_text = "\n".join(section_info["texts"])
            section_images = section_info.get("images", [])

            # 如果section过长，进行切分
            if len(section_text) > self.config.max_chunk_size:
                print(
                    f"  ✂️  章节 [{section_info['section']}] 过长 ({len(section_text)} 字符)，进行切分..."
                )
                sub_chunks = self._split_text_with_overlap(
                    text=section_text,
                    source_file=source_file,
                    section=section_info["section"],
                    images=section_images,  # 传递该 section 的图片
                    page_number=section_info["page"],
                )
                chunks.extend(sub_chunks)
            else:
                # 直接创建chunk
                if section_text.strip():
                    chunk = self._create_chunk(
                        text=section_text,
                        source_file=source_file,
                        section=section_info["section"],
                        images=section_images,
                        page_number=section_info["page"],
                    )
                    chunks.append(chunk)

        # 第三步：处理剩余未定位的图片（如果有）
        if remaining_images and chunks:
            chunks = self._distribute_images_to_chunks(chunks, remaining_images)

        print(f"  ✅ 生成 {len(chunks)} 个优化的文档块")
        return chunks

    def _is_custom_title(self, text: str) -> bool:
        """检查是否匹配自定义标题模式"""
        for pattern in self.config.title_patterns:
            if re.match(pattern, text.strip()):
                return True
        return False

    def _split_text_with_overlap(
        self,
        text: str,
        source_file: str,
        section: str,
        images: List[str] = None,
        page_number: int = 0,
    ) -> List[DocumentChunk]:
        """
        使用重叠策略切分长文本

        Args:
            text: 要切分的文本
            source_file: 源文件
            section: 章节名
            images: 关联图片
            page_number: 页码

        Returns:
            切分后的 chunk 列表
        """
        if images is None:
            images = []

        chunks = []

        # 第一步：保护代码块，避免被分割打断
        code_block_pattern = r"(```[\s\S]*?```)"
        code_blocks = []

        def protect_code_block(match):
            """将代码块替换为占位符"""
            code_blocks.append(match.group(0))
            return f"__CODE_BLOCK_{len(code_blocks) - 1}__"

        protected_text = re.sub(code_block_pattern, protect_code_block, text)

        def restore_code_blocks(t: str) -> str:
            """恢复代码块"""
            result = t
            for i, code_block in enumerate(code_blocks):
                result = result.replace(f"__CODE_BLOCK_{i}__", code_block)
            return result

        # 如果文本不超过最大长度，直接返回
        if len(text) <= self.config.max_chunk_size:
            chunk = self._create_chunk(
                text=text,
                source_file=source_file,
                section=section,
                images=images,
                page_number=page_number,
            )
            chunks.append(chunk)
            return chunks

        # 策略1: 尝试按段落切分（使用保护后的文本）
        if self.config.split_by_paragraph:
            paragraphs = protected_text.split("\n")
            paragraphs = [p.strip() for p in paragraphs if p.strip()]

            current_chunk_text = []
            current_length = 0

            for para in paragraphs:
                # 检查是否是代码块占位符（不应该被分割）
                is_code_placeholder = para.startswith(
                    "__CODE_BLOCK_"
                ) and para.endswith("__")
                para_len = len(para)

                # 如果是代码块占位符，获取实际长度
                actual_para_len = para_len
                if is_code_placeholder:
                    try:
                        idx = int(para.replace("__CODE_BLOCK_", "").replace("__", ""))
                        actual_para_len = (
                            len(code_blocks[idx])
                            if idx < len(code_blocks)
                            else para_len
                        )
                    except ValueError:
                        pass

                # 如果单个段落就超长，需要强制切分（但代码块作为整体保留）
                if (
                    actual_para_len > self.config.max_chunk_size
                    and not is_code_placeholder
                ):
                    # 保存当前累积的
                    if current_chunk_text:
                        chunk_text = restore_code_blocks("\n".join(current_chunk_text))
                        chunk = self._create_chunk(
                            text=chunk_text,
                            source_file=source_file,
                            section=section,
                            images=images if not chunks else [],
                            page_number=page_number,
                        )
                        chunks.append(chunk)
                        current_chunk_text = []
                        current_length = 0

                    # 强制切分超长段落
                    restored_para = restore_code_blocks(para)
                    sub_chunks = self._split_by_characters(
                        restored_para, source_file, section, [], page_number
                    )
                    chunks.extend(sub_chunks)

                # 如果加上这个段落会超长
                elif current_length + actual_para_len > self.config.max_chunk_size:
                    # 保存当前chunk
                    if current_chunk_text:
                        chunk_text = restore_code_blocks("\n".join(current_chunk_text))
                        chunk = self._create_chunk(
                            text=chunk_text,
                            source_file=source_file,
                            section=section,
                            images=images if not chunks else [],
                            page_number=page_number,
                        )
                        chunks.append(chunk)

                    # 添加重叠部分（代码块不参与重叠）
                    if (
                        self.config.chunk_overlap > 0
                        and current_chunk_text
                        and not is_code_placeholder
                    ):
                        overlap_text = current_chunk_text[-1]
                        current_chunk_text = [overlap_text, para]
                        current_length = len(overlap_text) + actual_para_len
                    else:
                        current_chunk_text = [para]
                        current_length = actual_para_len
                else:
                    current_chunk_text.append(para)
                    current_length += actual_para_len

            # 保存最后一个chunk
            if current_chunk_text:
                chunk_text = restore_code_blocks("\n".join(current_chunk_text))
                chunk = self._create_chunk(
                    text=chunk_text,
                    source_file=source_file,
                    section=section,
                    images=images if not chunks else [],
                    page_number=page_number,
                )
                chunks.append(chunk)

        else:
            # 策略2: 按固定字符数切分
            chunks = self._split_by_characters(
                text, source_file, section, images, page_number
            )

        return chunks

    def _split_by_characters(
        self,
        text: str,
        source_file: str,
        section: str,
        images: List[str],
        page_number: int,
    ) -> List[DocumentChunk]:
        """按固定字符数切分（带重叠）"""
        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = min(start + self.config.max_chunk_size, text_length)

            # 尽量在句号、问号、感叹号处切分
            if end < text_length:
                for sep in ["。", "！", "？", ".", "!", "?"]:
                    last_sep = text[start:end].rfind(sep)
                    if last_sep > self.config.max_chunk_size * 0.7:  # 至少达到70%长度
                        end = start + last_sep + 1
                        break

            chunk_text = text[start:end].strip()

            if chunk_text:
                chunk = self._create_chunk(
                    text=chunk_text,
                    source_file=source_file,
                    section=section,
                    images=images if not chunks else [],
                    page_number=page_number,
                )
                chunks.append(chunk)

            # 计算下一个起点（考虑重叠）
            if self.config.chunk_overlap > 0 and end < text_length:
                start = max(start + 1, end - self.config.chunk_overlap)
            else:
                start = end

        return chunks

    def _distribute_images_to_chunks(
        self, chunks: List[DocumentChunk], images: List[str]
    ) -> List[DocumentChunk]:
        """
        智能分配图片到各个chunk

        策略：根据chunk数量均匀分配，或者集中分配到第一个chunk
        """
        if not images or not chunks:
            return chunks

        if self.config.distribute_images_evenly:
            # 均匀分配
            images_per_chunk = max(1, len(images) // len(chunks))
            for i, chunk in enumerate(chunks):
                start_idx = i * images_per_chunk
                end_idx = start_idx + images_per_chunk
                if i == len(chunks) - 1:  # 最后一个chunk获取剩余所有图片
                    end_idx = len(images)
                chunk.images = images[start_idx:end_idx]
                chunk.metadata["image_count"] = len(chunk.images)
                chunk.metadata["has_images"] = len(chunk.images) > 0
        else:
            # 集中分配到第一个chunk
            chunks[0].images = images
            chunks[0].metadata["image_count"] = len(images)
            chunks[0].metadata["has_images"] = True

        return chunks

    def _create_chunk(
        self,
        text: str,
        source_file: str,
        section: str,
        images: List[str],
        page_number: int,
    ) -> DocumentChunk:
        """创建文档块"""
        chunk_id = hashlib.md5(
            f"{source_file}:{section}:{text[:100]}".encode()
        ).hexdigest()[:12]

        return DocumentChunk(
            chunk_id=chunk_id,
            text=text,
            source_file=source_file,
            section=section,
            page_number=page_number,
            images=images,
            metadata={
                "has_images": len(images) > 0,
                "image_count": len(images),
                "text_length": len(text),
            },
        )

    def _get_page_number(self, element) -> int:
        """获取元素所在页码"""
        if hasattr(element, "metadata") and hasattr(element.metadata, "page_number"):
            return element.metadata.page_number or 0
        return 0
